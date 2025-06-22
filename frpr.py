from docx import Document
from docx.shared import Inches

# Création du document
doc = Document()
doc.add_heading("Projet : Plateforme Citoyenne de Signalement pour Cotonou", 0)

# Objectif
doc.add_heading("🎯 Objectif", level=1)
doc.add_paragraph(
    "Permettre aux citoyens de signaler des problèmes urbains (ordures non ramassées, trous dans les routes, "
    "lampadaires défectueux, inondations, etc.), tout en informant la mairie, les services publics ou les "
    "associations locales, dans un esprit de collaboration civique."
)

# Fonctionnalités principales
doc.add_heading("🧩 Fonctionnalités principales", level=1)

doc.add_heading("🧭 Signalement rapide", level=2)
doc.add_paragraph("- Catégories : ordures, voirie, éclairage, insécurité, inondation, bruit, transport, autres.")
doc.add_paragraph("- Ajout de photo et géolocalisation automatique.")
doc.add_paragraph("- Description courte.")
doc.add_paragraph("- Suivi anonyme ou avec compte.")

doc.add_heading("📍 Carte interactive", level=2)
doc.add_paragraph("- Affichage des signalements publics par quartier.")
doc.add_paragraph("- Filtres (problèmes résolus, urgents, en attente).")
doc.add_paragraph("- Vue satellite ou plan via OpenStreetMap.")

doc.add_heading("👥 Profil citoyen ou association", level=2)
doc.add_paragraph("- Historique de signalements.")
doc.add_paragraph("- Système de confiance : les signalements confirmés gagnent du poids.")
doc.add_paragraph("- Option de partage sur WhatsApp ou Facebook.")

doc.add_heading("🏛️ Interface pour les mairies ou services techniques", level=2)
doc.add_paragraph("- Tableau de bord des signalements.")
doc.add_paragraph("- Classement par type, urgence, localisation.")
doc.add_paragraph("- Statistiques (par quartier, taux de résolution, temps moyen de traitement).")

# Version mobile
doc.add_heading("📱 Version mobile-first", level=1)
doc.add_paragraph(
    "Interface légère et rapide, conçue pour smartphones Android. Fonctionne même en connexion limitée "
    "(mode offline temporaire). Progressive Web App (PWA) installable."
)

# Problématiques locales
doc.add_heading("🔒 Problématiques locales à anticiper", level=1)
doc.add_paragraph("Tableau des enjeux et solutions :")
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Enjeu'
hdr_cells[1].text = 'Solution adaptée'

issues = [
    ("Faible confiance dans les autorités", "Option d’anonymat, publication publique par défaut"),
    ("Faible connectivité internet", "Mode hors-ligne + envoi différé des signalements"),
    ("Manque de réactivité des services publics", "Implication d’associations ou d’acteurs alternatifs"),
    ("Alphabétisation variée", "Icônes explicites + interface traduite en français et fon/yoruba"),
    ("Appareils peu puissants", "Interface optimisée (React + Tailwind léger)"),
]

for issue, solution in issues:
    row_cells = table.add_row().cells
    row_cells[0].text = issue
    row_cells[1].text = solution

# Stack technique
doc.add_heading("🔧 Stack technique adaptée", level=1)
stack = {
    "Frontend": "React + TailwindCSS",
    "Backend": "Python FastAPI",
    "Base de données": "PostgreSQL (avec PostGIS pour géodonnées)",
    "Cartographie": "OpenStreetMap + Leaflet.js",
    "Authentification": "Firebase Auth ou JWT",
    "Déploiement": "Docker, hébergement local ou VPS (HostAfrica, etc.)"
}

for comp, tech in stack.items():
    doc.add_paragraph(f"{comp} : {tech}")

# Plan de déploiement
doc.add_heading("🚀 Plan de déploiement progressif (MVP)", level=1)
doc.add_paragraph("1. MVP quartier-pilote (ex. Zogbo, Akpakpa ou Fidjrossè)")
doc.add_paragraph("2. Ajout d’un tableau de bord basique pour les associations partenaires")
doc.add_paragraph("3. Campagne de sensibilisation : flyers, Facebook, radios locales")
doc.add_paragraph("4. Négociation avec la mairie pour tester l’usage interne")
doc.add_paragraph("5. Extension progressive à d’autres arrondissements")

# Acteurs à impliquer
doc.add_heading("🤝 Acteurs à impliquer", level=1)
actors = [
    "Mairies d’arrondissement",
    "ONG locales (WILDAF, Social Watch, Terre des Hommes…)",
    "Universités / étudiants pour la promotion et les tests terrain",
    "Radios communautaires pour relayer les signalements"
]
for actor in actors:
    doc.add_paragraph(f"- {actor}")

# Sauvegarder le fichier
file_path = "Plateforme_Signalement_Cotonou.docx"
doc.save(file_path)